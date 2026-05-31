"""
Gera Trabalho-2-PataAmiga.docx a partir do template oficial usando python-docx.

Estrategia: carrega o template (`Trabalho 2 - Template - Documentacao de Projeto - 20 pontos.docx`)
para herdar estilos (Heading1/2/3, Title, paginacao A4 com margens), limpa o body
e reconstroi o documento conforme a estrutura exigida.

Uso:
    python generate_docx.py

Requer:
    pip install python-docx
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Sequence

# Console do Windows usa cp1252 por default; o nome do template tem caracteres
# acentuados que estouram em print(). Forca UTF-8 quando disponivel.
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ============================================================================
#  Paths
# ============================================================================

SCRIPT_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
IMG_DIR      = PROJECT_ROOT / "docs" / "diagrams" / "png"
OUTPUT       = SCRIPT_DIR / "Trabalho-2-PataAmiga.docx"


def find_template() -> Path:
    candidates: list[Path] = []
    candidates += list(PROJECT_ROOT.glob("Trabalho*Template*.docx"))
    candidates += list((PROJECT_ROOT / "templates-examples").glob("Trabalho*Template*.docx"))
    if candidates:
        return candidates[0]
    # Fallback: o template oficial e local/gitignored e pode estar ausente. O .docx
    # gerado anteriormente herdou os mesmos estilos (Title, Heading 1-4, A4/margens),
    # entao serve de base. clear_body() limpa o conteudo e o documento e reconstruido.
    if OUTPUT.exists():
        print(
            f"AVISO: template oficial nao encontrado; reutilizando estilos de "
            f"{OUTPUT.name} (saida anterior)."
        )
        return OUTPUT
    sys.exit(
        "ERRO: template `.docx` nao encontrado. Esperado um arquivo "
        "`Trabalho*Template*.docx` na raiz do projeto ou em `templates-examples/`, "
        "ou um `Trabalho-2-PataAmiga.docx` gerado anteriormente nesta pasta."
    )


# ============================================================================
#  Helpers de baixo nivel (OOXML)
# ============================================================================

def clear_body(doc) -> None:
    """Remove paragrafos e tabelas, mantendo `<w:sectPr>` (paginacao/margens).

    Tambem descarta os relacionamentos de imagem: como o corpo foi esvaziado, toda
    imagem herdada da base fica orfa. Sem isso, reusar a saida anterior como base
    (ver find_template) acumularia midias mortas a cada regeracao. add_image()
    recria os relacionamentos necessarios.
    """
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)

    rels = doc.part.rels
    for rId in [rId for rId, rel in rels.items() if "image" in rel.reltype]:
        del rels[rId]


def add_top_border(paragraph) -> None:
    """Linha horizontal grossa no topo do paragrafo (decorativa da capa)."""
    pPr   = paragraph._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    top   = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "36")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "000000")
    pBdr.append(top)
    pPr.append(pBdr)


def insert_toc(doc, title: str = "Tabela de Conteudo") -> None:
    """Insere o titulo do sumario + um campo TOC do Word (\\o "1-3" \\h \\z \\u)."""
    # Titulo do sumario (sem entrar na propria TOC)
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.bold = True

    # Instrucoes ao Word para gerar o sumario
    p = doc.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    placeholder = OxmlElement("w:t")
    placeholder.text = "Sumario sera populado pelo Word (botao direito > Atualizar Campo, ou F9)."
    run._r.append(placeholder)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


# ============================================================================
#  Helpers de alto nivel (conteudo)
# ============================================================================

def add_page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_cover_line(
    doc,
    text: str,
    size_pt: float,
    *,
    bold: bool = True,
    italic: bool = False,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    space_after_pt: float = 12,
    space_before_pt: float = 0,
):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after  = Pt(space_after_pt)
    pf.space_before = Pt(space_before_pt)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    return p


def add_para(doc, text: str, *, justify: bool = True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullets(doc, items: Sequence[str]) -> None:
    """O template nao traz estilo 'List Bullet'; emulamos com bullet char + recuo."""
    for item in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent       = Cm(0.75)
        pf.first_line_indent = Cm(-0.5)
        pf.space_after       = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run("•\t" + item)


def _set_cell_bold(cell, bold: bool = True) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold


def _shade_cell(cell, hex_fill: str = "DDDDDD") -> None:
    """Preenche o cell com cor de fundo (cabecalho)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    tcPr.append(shd)


def _set_table_borders(tbl, sz: str = "4", color: str = "888888") -> None:
    """Adiciona bordas a todas as celulas (top/left/bottom/right/insideH/insideV)."""
    tblPr  = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def add_table(
    doc,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
) -> None:
    cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=cols)
    _set_table_borders(tbl)
    # Cabecalho
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _set_cell_bold(hdr_cells[i], True)
        _shade_cell(hdr_cells[i], "E7E7E7")
    # Corpo
    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    doc.add_paragraph()


def add_kv_table(doc, pairs: Sequence[tuple[str, str]]) -> None:
    """Tabela 2 colunas com cabecalho 'Campo / Descricao' e bold na esquerda."""
    tbl = doc.add_table(rows=1 + len(pairs), cols=2)
    _set_table_borders(tbl)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Descricao"
    _set_cell_bold(hdr[0])
    _set_cell_bold(hdr[1])
    _shade_cell(hdr[0], "E7E7E7")
    _shade_cell(hdr[1], "E7E7E7")
    for i, (k, v) in enumerate(pairs):
        row = tbl.rows[i + 1].cells
        row[0].text = k
        _set_cell_bold(row[0])
        row[1].text = v
    doc.add_paragraph()


def add_image(doc, filename: str, caption: str, *, width_cm: float = 15) -> None:
    img_path = IMG_DIR / filename
    if not img_path.exists():
        sys.exit(f"ERRO: imagem nao encontrada: {img_path}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cap.add_run(caption)
    crun.italic = True
    crun.font.size = Pt(10)
    doc.add_paragraph()


def add_code(doc, code_text: str) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def h1(doc, text: str):  return doc.add_heading(text, level=1)
def h2(doc, text: str):  return doc.add_heading(text, level=2)
def h3(doc, text: str):  return doc.add_heading(text, level=3)
def h4(doc, text: str):  return doc.add_heading(text, level=4)


# ============================================================================
#  Conteudo do documento
# ============================================================================

def build_cover(doc) -> None:
    # Linha decorativa do topo
    p_top = doc.add_paragraph()
    add_top_border(p_top)

    add_cover_line(doc, "Documentacao de Projeto",   32, space_after_pt=36)
    add_cover_line(doc, "para o sistema PataAmiga",  24, italic=True, space_after_pt=40)
    add_cover_line(doc, "Versao 1.0",                18, italic=True, space_after_pt=80)
    add_cover_line(
        doc,
        "Projeto de sistema elaborado pelo aluno Renato Douglas "
        "como parte da disciplina Projeto de Software.",
        14, bold=False, space_after_pt=40,
    )
    add_cover_line(doc, "23 de maio de 2026", 14, bold=False, italic=True)
    add_page_break(doc)


def build_historico(doc) -> None:
    h1(doc, "Historico de Revisoes")
    add_table(doc,
        headers=["Nome", "Data", "Razoes para Mudanca", "Versao"],
        rows=[[
            "Renato Douglas",
            "23/05/2026",
            "Versao inicial - entrega do Trabalho 2 da disciplina Projeto de Software",
            "1.0",
        ]],
    )
    add_page_break(doc)


def build_toc(doc) -> None:
    insert_toc(doc)
    add_page_break(doc)


# ----- 1. Introducao -------------------------------------------------------

def build_intro(doc) -> None:
    h1(doc, "1. Introducao")
    add_para(doc,
        "Este documento agrega: (1) a elaboracao e revisao de modelos de dominio "
        "e (2) modelos de projeto para o sistema PataAmiga. A referencia principal "
        "para a descricao geral do problema, dominio e requisitos do sistema e o "
        "documento de especificacao que descreve a visao de dominio do sistema."
    )

    h2(doc, "1.1. Sobre o PataAmiga")
    add_para(doc,
        "O PataAmiga e uma plataforma web desenvolvida para centralizar e otimizar "
        "a gestao de ONGs dedicadas ao resgate e adocao de animais. O sistema "
        "acompanha o ciclo de vida completo de cada animal - desde o resgate inicial, "
        "passando pelos cuidados veterinarios, ate a adocao responsavel - integrando "
        "todos os atores envolvidos (adotantes, voluntarios, veterinarios, "
        "coordenadores e doadores) em uma unica solucao."
    )
    add_para(doc,
        "ONGs do setor frequentemente operam com processos manuais e descentralizados "
        "(planilhas, grupos de mensagens, cadernos fisicos), o que dificulta o "
        "controle, compromete a rastreabilidade e limita a capacidade de atendimento. "
        "O PataAmiga resolve esse problema oferecendo uma plataforma centralizada que:"
    )
    add_bullets(doc, [
        "Registra e acompanha cada animal desde o resgate ate a adocao;",
        "Organiza o historico veterinario completo de cada animal;",
        "Gerencia o processo de adocao com etapas formalizadas;",
        "Conecta adotantes a animais disponiveis com busca por perfil;",
        "Permite que doadores acompanhem o impacto de suas contribuicoes;",
        "Fornece ao coordenador da ONG visibilidade operacional e relatorios gerenciais.",
    ])

    h2(doc, "1.2. Stack tecnologica adotada")
    add_table(doc,
        headers=["Camada", "Tecnologia"],
        rows=[
            ["Frontend",                     "Angular 17 + TypeScript"],
            ["Backend",                      "Java 21 + Spring Boot 3.x"],
            ["Banco de Dados",               "PostgreSQL 16"],
            ["ORM",                          "Hibernate / JPA"],
            ["Mapeamento DTO <-> Entity",    "MapStruct"],
            ["Autenticacao",                 "JWT + Spring Security"],
            ["Containerizacao",              "Docker + Docker Compose"],
        ],
    )


# ----- 2. Modelos de Usuario e Requisitos ----------------------------------

def build_secao2(doc) -> None:
    h1(doc, "2. Modelos de Usuario e Requisitos")

    # 2.1
    h2(doc, "2.1. Descricao de Atores")
    add_para(doc,
        "O PataAmiga possui cinco atores que interagem com o sistema, agrupados em "
        "externos (clientes finais) e internos (operadores da ONG). A tabela a seguir "
        "resume papel, tipo e responsabilidades principais de cada um."
    )
    add_table(doc,
        headers=["Ator", "Tipo", "Responsabilidades"],
        rows=[
            ["Adotante",            "Externo / Primario",       "Buscar animais disponiveis, submeter solicitacao de adocao, acompanhar o processo e registrar pos-adocao"],
            ["Voluntario",          "Interno / Operacional",    "Registrar animais resgatados, agendar atendimentos veterinarios, atualizar status e apoiar eventos"],
            ["Veterinario",         "Interno / Especialista",   "Emitir laudos clinicos, registrar prontuarios, aplicar vacinas e procedimentos, liberar animais para adocao"],
            ["Coordenador da ONG",  "Interno / Administrativo", "Gerenciar equipe, aprovar ou rejeitar adocoes, controlar recursos, emitir relatorios gerenciais"],
            ["Doador",              "Externo / Suporte",        "Realizar doacoes pontuais ou recorrentes, acompanhar uso dos recursos e visualizar relatorios de impacto"],
        ],
    )
    add_para(doc,
        "Decisao de modelagem (Abordagem C). Embora os atores sejam apresentados de "
        "forma separada na visao UML, no modelo de dados existe uma unica entidade "
        "Usuario central com cinco perfis 1:1 opcionais (PerfilAdotante, PerfilDoador, "
        "PerfilVoluntario, PerfilVeterinario, PerfilCoordenador). Uma mesma pessoa pode "
        "acumular qualquer combinacao de papeis. O detalhamento aparece na Secao 3.3 "
        "(Diagrama de Classes) e na Secao 4 (Modelos de Dados)."
    )

    # 2.2
    h2(doc, "2.2. Modelo de Casos de Uso")
    add_para(doc,
        "Os 14 casos de uso do sistema estao agrupados em 5 pacotes funcionais. "
        "Cada caso de uso recebe um ID UC-NN usado como referencia cruzada nas "
        "demais secoes."
    )
    add_table(doc,
        headers=["ID", "Caso de Uso", "Ator(es) Primario(s)", "Pacote"],
        rows=[
            ["UC-01", "Solicitar Adocao",                       "Adotante",            "Adocao"],
            ["UC-02", "Registrar Animal Resgatado",             "Voluntario",          "Gestao de Animais"],
            ["UC-03", "Registrar Atendimento Veterinario",      "Veterinario",         "Atendimento Veterinario"],
            ["UC-04", "Buscar Animal Disponivel",               "Adotante",            "Adocao"],
            ["UC-05", "Acompanhar Processo de Adocao",          "Adotante",            "Adocao"],
            ["UC-06", "Registrar Pos-Adocao",                   "Adotante",            "Adocao"],
            ["UC-07", "Atualizar Status do Animal",             "Voluntario",          "Gestao de Animais"],
            ["UC-08", "Agendar Atendimento Veterinario",        "Voluntario",          "Gestao de Animais"],
            ["UC-09", "Emitir Laudo de Liberacao para Adocao",  "Veterinario",         "Atendimento Veterinario"],
            ["UC-10", "Aprovar/Rejeitar Adocao",                "Coordenador da ONG",  "Adocao"],
            ["UC-11", "Gerenciar Voluntarios",                  "Coordenador da ONG",  "Administracao"],
            ["UC-12", "Emitir Relatorios Gerenciais",           "Coordenador da ONG",  "Administracao"],
            ["UC-13", "Realizar Doacao",                        "Doador",              "Doacoes"],
            ["UC-14", "Acompanhar Destinacao das Doacoes",      "Doador",              "Doacoes"],
        ],
    )

    h3(doc, "2.2.1. Relacionamentos entre Casos de Uso")
    add_table(doc,
        headers=["Relacionamento", "Descricao"],
        rows=[
            ["UC-03 «include» UC-07", "Todo atendimento veterinario registra uma transicao de status (ex.: RESGATADO -> EM_TRATAMENTO)."],
            ["UC-09 «include» UC-07", "A emissao do laudo de liberacao move o animal para DISPONIVEL."],
            ["UC-10 «include» UC-07", "A aprovacao move para EM_PROCESSO_ADOCAO/ADOTADO; a rejeicao mantem em DISPONIVEL."],
        ],
    )

    add_image(doc, "01-caso-de-uso.png", "Figura 1 - Diagrama de Casos de Uso do PataAmiga")

    # 2.3
    h2(doc, "2.3. Diagrama de Sequencia do Sistema")
    add_para(doc,
        "Apresenta-se a seguir o Diagrama de Sequencia do Sistema (SSD) do PataAmiga, "
        "em visao caixa-preta e ponta a ponta. Os cinco atores interagem com :Sistema "
        "por meio das operacoes de sistema, cobrindo o ciclo de vida do animal (resgate "
        "-> tratamento -> liberacao -> adocao -> pos-adocao) e o fluxo paralelo de "
        "doacoes. Por ser caixa-preta, o diagrama omite as camadas internas; a "
        "realizacao interna (caixa-branca) de cada operacao e detalhada nos diagramas "
        "de sequencia da Secao 3.4 (Figuras 8 a 21)."
    )
    add_table(doc,
        headers=["Operacao de Sistema", "Ator", "Caso de Uso"],
        rows=[
            ["registrarAnimalResgatado(dados)",                       "Voluntario",  "UC-02"],
            ["registrarAtendimento(idAnimal, dados)",                 "Veterinario", "UC-03 «include» UC-07"],
            ["emitirLaudoLiberacao(idAnimal, parecer)",              "Veterinario", "UC-09 «include» UC-07"],
            ["buscarAnimaisDisponiveis(especie, porte, sexo)",       "Adotante",    "UC-04"],
            ["submeterSolicitacaoAdocao(idAnimal, formularioInteresse)", "Adotante", "UC-01"],
            ["aprovarAdocao(idAdocao) / concluirAdocao(idAdocao)",   "Coordenador", "UC-10 «include» UC-07"],
            ["registrarPosAdocao(idAdocao, relato)",                 "Adotante",    "UC-06"],
            ["realizarDoacao(valor, tipo, metodo)",                  "Doador",      "UC-13"],
        ],
    )
    add_image(doc, "02-ssd-sistema.png",
              "Figura 2 - Diagrama de Sequencia do Sistema do PataAmiga (caixa-preta, ponta-a-ponta)",
              width_cm=12)


# ----- 3. Modelos de Projeto -----------------------------------------------

def build_secao3(doc) -> None:
    h1(doc, "3. Modelos de Projeto")

    # 3.1
    h2(doc, "3.1. Arquitetura")
    add_para(doc,
        "O PataAmiga adota uma arquitetura monolitica modular em camadas, adequada "
        "para o porte de uma ONG de pequeno e medio porte. A arquitetura e descrita "
        "usando o C4 Model em tres niveis (Contexto, Containers e Componentes)."
    )

    h3(doc, "3.1.1. Organizacao de pacotes do backend")
    add_para(doc,
        "O codigo do backend segue a organizacao tipica de uma aplicacao Spring Boot "
        "baseada em camadas, sob a raiz br.pucminas.pataamiga:"
    )
    add_code(doc,
        "controller/          @RestController, mapeamento de rotas REST\n"
        "service/             @Service, regras de negocio\n"
        "repository/          interfaces Spring Data JPA (@Repository)\n"
        "model/               entidades JPA (@Entity, @Table)\n"
        "dto/\n"
        "  | request/         payloads de entrada  (AnimalRequestDTO...)\n"
        "  | response/        payloads de saida   (AnimalResponseDTO...)\n"
        "mapper/              conversao DTO <-> Entity via MapStruct (@Mapper)\n"
        "exception/           exceptions customizadas + @RestControllerAdvice\n"
        "config/              SecurityConfig, OpenApiConfig, CorsConfig\n"
        "security/            JwtFilter, JwtUtil, UserDetailsServiceImpl"
    )
    add_para(doc,
        "Padroes adotados: Repository Pattern, Service Layer, DTO (request/response), "
        "Mapper (MapStruct), Global Exception Handling, Spring Security com RBAC por "
        "perfil de ator."
    )

    h3(doc, "3.1.2. C4 Nivel 1 - Contexto")
    add_image(doc, "03-c4-contexto.png",   "Figura 3 - C4 Nivel 1: Diagrama de Contexto do PataAmiga")

    h3(doc, "3.1.3. C4 Nivel 2 - Containers")
    add_image(doc, "04-c4-containers.png", "Figura 4 - C4 Nivel 2: Diagrama de Containers do PataAmiga")

    # 3.2
    h2(doc, "3.2. Diagrama de Componentes e Implantacao")
    add_para(doc,
        "O diagrama de componentes corresponde ao C4 Nivel 3, detalhando os "
        "componentes internos do backend Spring Boot. O diagrama de implantacao mostra "
        "onde cada container e alocado na infraestrutura Docker."
    )

    h3(doc, "3.2.1. Diagrama de Componentes (C4 Nivel 3)")
    add_image(doc, "05-c4-componentes.png", "Figura 5 - C4 Nivel 3: Componentes do Backend PataAmiga")

    h3(doc, "3.2.2. Diagrama de Implantacao")
    add_para(doc,
        "A implantacao e orquestrada por Docker Compose (arquivo docker-compose.yml na "
        "raiz do projeto). Tres containers compartilham a rede bridge pataamiga-net em "
        "um host Linux com Docker Engine, e dois sistemas externos sao consumidos pela API."
    )
    add_table(doc,
        headers=["Container", "Imagem Docker", "Portas (host:container)", "Responsabilidade", "Volume"],
        rows=[
            ["pataamiga-frontend", "nginx:alpine",            "4200:80",   "Servir o build estatico do SPA Angular 17",                       "-"],
            ["pataamiga-backend",  "eclipse-temurin:21-jre",  "8080:8080", "API REST Spring Boot 3.x (/api/v1/**) com JWT e RBAC",            "-"],
            ["pataamiga-db",       "postgres:16-alpine",      "5432:5432", "Banco PostgreSQL com schema pataamiga",                           "pataamiga-pgdata"],
        ],
    )
    add_table(doc,
        headers=["Sistema externo", "Protocolo", "Uso"],
        rows=[
            ["Servico de E-mail (SMTP)", "SMTP/TLS",   "Notificacoes de adocao e confirmacoes de doacao"],
            ["Gateway de Pagamento",     "HTTPS/REST", "Processamento de doacoes via Pix e cartao"],
        ],
    )
    add_para(doc,
        "Ordem de inicializacao: pataamiga-db (com healthcheck) -> pataamiga-backend "
        "(depends_on: db) -> pataamiga-frontend (depends_on: backend)."
    )
    add_image(doc, "06-implantacao.png", "Figura 6 - Diagrama de Implantacao do PataAmiga (Docker Compose)")

    # 3.3
    h2(doc, "3.3. Diagrama de Classes")
    add_para(doc,
        "O modelo de dominio adota a Abordagem C para usuarios: uma entidade Usuario "
        "central concentra os dados pessoais (nome, e-mail, CPF, telefone, endereco, "
        "senha), e cinco perfis especializados - PerfilAdotante, PerfilDoador, "
        "PerfilVoluntario, PerfilVeterinario e PerfilCoordenador - estendem Usuario por "
        "composicao 1:1 opcional. Uma mesma pessoa pode acumular qualquer combinacao de "
        "papeis sem duplicar dados de cadastro."
    )

    h3(doc, "3.3.1. Principais classes do dominio")
    add_table(doc,
        headers=["Classe", "Tipo", "Papel no modelo"],
        rows=[
            ["Usuario",                                      "Entidade central",            "Identidade comum a todos os atores; responsavel por autenticacao"],
            ["Perfil (abstrata)",                            "Superclasse",                 "Define ativo, dataAtivacao e operacoes genericas dos perfis"],
            ["PerfilAdotante / Doador / Voluntario / Veterinario / Coordenador",
                                                             "Perfis 1:1",                  "Especializam Usuario com atributos e operacoes especificas de cada papel"],
            ["Animal",                                       "Entidade central do negocio", "Estado controlado pelo enum StatusAnimal; relaciona-se com prontuarios e adocoes"],
            ["RegistroVeterinario",                          "Entidade",                    "Cada atendimento veterinario registrado; liberaParaAdocao aciona transicao de status"],
            ["Adocao",                                       "Entidade",                    "Processo de adocao com ciclo de vida EM_ANALISE -> ... -> CONCLUIDA/REJEITADA"],
            ["Doacao",                                       "Entidade",                    "Aporte pontual ou recorrente associado a um PerfilDoador"],
            ["Relatorio",                                    "Entidade",                    "Saida agregada gerada por um PerfilCoordenador"],
            ["StatusAnimal, StatusAdocao, TipoDoacao, TipoRelatorio, TipoAtendimento, Especie, Sexo, Porte",
                                                             "Enums",                       "Estados e classificadores do dominio"],
        ],
    )
    add_image(doc, "07-classes.png", "Figura 7 - Diagrama de Classes do Modelo de Dominio do PataAmiga")

    # 3.4
    h2(doc, "3.4. Diagramas de Sequencia")
    add_para(doc,
        "Esta secao apresenta a realizacao interna (caixa-branca) das operacoes de "
        "sistema definidas no Diagrama de Sequencia do Sistema da Secao 2.3. Diferente "
        "do SSD - que trata o sistema como uma caixa-preta (ator <-> :Sistema) -, cada "
        "diagrama de sequencia abre o :Sistema e detalha a colaboracao entre as camadas "
        "da arquitetura - SPA Angular -> JwtFilter -> Controller -> Service -> Repository "
        "-> PostgreSQL -, reutilizando os mesmos componentes do C4 Nivel 3 (Figura 5). "
        "Apresenta-se a realizacao individual de cada um dos 14 casos de uso (UC-01 a UC-14)."
    )

    h3(doc, "3.4.1. Realizacao por caso de uso")
    add_para(doc,
        "A tabela a seguir mapeia cada caso de uso aos componentes que participam da sua "
        "realizacao (Controllers, Services e Repositories), seguida pelos respectivos "
        "diagramas de sequencia caixa-branca."
    )
    add_table(doc,
        headers=["Fig.", "Caso de Uso", "Controllers", "Services", "Repositories"],
        rows=[
            ["8",  "UC-01 Solicitar Adocao",                  "AnimalController, AdocaoController",      "AnimalService, AdocaoService",      "AnimalRepository, AdocaoRepository"],
            ["9",  "UC-02 Registrar Animal Resgatado",        "AnimalController",                        "AnimalService",                     "AnimalRepository, RegistroVetRepository"],
            ["10", "UC-03 Registrar Atendimento Veterinario", "AnimalController, RegistroVetController",  "AnimalService, RegistroVetService", "AnimalRepository, RegistroVetRepository"],
            ["11", "UC-04 Buscar Animal Disponivel",          "AnimalController",                        "AnimalService",                     "AnimalRepository"],
            ["12", "UC-05 Acompanhar Processo de Adocao",     "AdocaoController",                        "AdocaoService",                     "AdocaoRepository"],
            ["13", "UC-06 Registrar Pos-Adocao",              "AdocaoController",                        "AdocaoService, AnimalService",      "AdocaoRepository, AnimalRepository"],
            ["14", "UC-07 Atualizar Status do Animal",        "AnimalController",                        "AnimalService",                     "AnimalRepository"],
            ["15", "UC-08 Agendar Atendimento Veterinario",   "RegistroVetController",                   "RegistroVetService",                "AnimalRepository, RegistroVetRepository"],
            ["16", "UC-09 Emitir Laudo de Liberacao",         "RegistroVetController",                   "RegistroVetService, AnimalService", "AnimalRepository, RegistroVetRepository"],
            ["17", "UC-10 Aprovar/Rejeitar Adocao",           "AdocaoController",                        "AdocaoService, AnimalService",      "AdocaoRepository, AnimalRepository"],
            ["18", "UC-11 Gerenciar Voluntarios",             "UsuarioController",                       "UsuarioService",                    "UsuarioRepository"],
            ["19", "UC-12 Emitir Relatorios Gerenciais",      "DoacaoController",                        "DoacaoService",                     "AdocaoRepository, DoacaoRepository, RelatorioRepository"],
            ["20", "UC-13 Realizar Doacao",                   "DoacaoController",                        "DoacaoService",                     "DoacaoRepository"],
            ["21", "UC-14 Acompanhar Destinacao das Doacoes", "DoacaoController",                        "DoacaoService",                     "DoacaoRepository"],
        ],
    )
    add_image(doc, "08-seq-solicitar-adocao.png",        "Figura 8 - Realizacao UC-01 Solicitar Adocao")
    add_image(doc, "09-seq-registrar-animal.png",        "Figura 9 - Realizacao UC-02 Registrar Animal Resgatado")
    add_image(doc, "10-seq-registrar-atendimento.png",   "Figura 10 - Realizacao UC-03 Registrar Atendimento Veterinario")
    add_image(doc, "11-seq-buscar-animal.png",           "Figura 11 - Realizacao UC-04 Buscar Animal Disponivel")
    add_image(doc, "12-seq-acompanhar-adocao.png",       "Figura 12 - Realizacao UC-05 Acompanhar Processo de Adocao")
    add_image(doc, "13-seq-registrar-pos-adocao.png",    "Figura 13 - Realizacao UC-06 Registrar Pos-Adocao")
    add_image(doc, "14-seq-atualizar-status.png",        "Figura 14 - Realizacao UC-07 Atualizar Status do Animal")
    add_image(doc, "15-seq-agendar-atendimento.png",     "Figura 15 - Realizacao UC-08 Agendar Atendimento Veterinario")
    add_image(doc, "16-seq-emitir-laudo.png",            "Figura 16 - Realizacao UC-09 Emitir Laudo de Liberacao para Adocao")
    add_image(doc, "17-seq-aprovar-rejeitar-adocao.png", "Figura 17 - Realizacao UC-10 Aprovar/Rejeitar Adocao")
    add_image(doc, "18-seq-gerenciar-voluntarios.png",   "Figura 18 - Realizacao UC-11 Gerenciar Voluntarios")
    add_image(doc, "19-seq-emitir-relatorios.png",       "Figura 19 - Realizacao UC-12 Emitir Relatorios Gerenciais")
    add_image(doc, "20-seq-realizar-doacao.png",         "Figura 20 - Realizacao UC-13 Realizar Doacao")
    add_image(doc, "21-seq-acompanhar-destinacao.png",   "Figura 21 - Realizacao UC-14 Acompanhar Destinacao das Doacoes")
    add_para(doc, "Observacoes de modelagem:")
    add_bullets(doc, [
        "A autenticacao e representada pelo passo validarToken() em JwtFilter, "
        "exigido em toda requisicao autenticada.",
        "O UC-03 ilustra explicitamente o relacionamento «include» UC-07 "
        "(Atualizar Status do Animal) - RegistroVetService delega a "
        "AnimalService.atualizarStatus(...) apos registrar o atendimento; o mesmo "
        "ocorre em UC-09 e UC-10.",
        "Validacoes de regra de negocio (ex.: Animal.status == DISPONIVEL antes de "
        "criar Adocao) aparecem como auto-mensagens ou notas no respectivo Service.",
        "O Diagrama de Sequencia do Sistema (Figura 2) apresenta a visao caixa-preta "
        "ponta-a-ponta; cada operacao de sistema e refinada (caixa-branca) no diagrama "
        "de sequencia do respectivo caso de uso (Figuras 8 a 21).",
    ])

    # 3.5
    h2(doc, "3.5. Diagramas de Comunicacao")
    add_para(doc,
        "Os diagramas de comunicacao apresentam a mesma colaboracao dos diagramas de "
        "sequencia da Secao 3.4, porem vistos como rede de objetos com numeracao "
        "explicita de mensagens em vez de eixo temporal vertical."
    )
    add_table(doc,
        headers=["Aspecto", "Sequencia (Secao 3.4)", "Comunicacao (Secao 3.5)"],
        rows=[
            ["Foco",      "Ordem temporal das mensagens", "Vinculos estruturais entre objetos"],
            ["Layout",    "Linhas de vida verticais",     "Objetos em rede"],
            ["Ordenacao", "Posicao vertical",             "Numeracao explicita (1:, 2:, ...)"],
        ],
    )
    add_para(doc,
        "Cada par de objetos e conectado por uma aresta sem direcao que carrega todas "
        "as mensagens trocadas entre eles - o sentido de cada mensagem e indicado pelos "
        "simbolos de direcao conforme o padrao UML. Os repositorios AnimalRepository e "
        "RegistroVetRepository aparecem lado a lado para refletir que ambos colaboram "
        "diretamente com o PostgreSQL. Para o UC-03, o diagrama modela apenas a operacao "
        "principal (registrarAtendimento) incluindo a transicao «include» UC-07 "
        "(VetSvc -> AnimSvc -> AnimRepo -> DB) - o fluxo alternativo de emissao de "
        "laudo e coberto no diagrama de sequencia (Figura 16)."
    )
    add_image(doc, "22-com-solicitar-adocao.png",     "Figura 22 - Comunicacao UC-01 Solicitar Adocao")
    add_image(doc, "23-com-registrar-animal.png",     "Figura 23 - Comunicacao UC-02 Registrar Animal Resgatado")
    add_image(doc, "24-com-registrar-atendimento.png","Figura 24 - Comunicacao UC-03 Registrar Atendimento Veterinario")

    # 3.6
    h2(doc, "3.6. Diagramas de Estados")
    add_para(doc,
        "Ciclo de vida de um animal dentro da plataforma, do resgate a adocao. Cada "
        "transicao e acionada por uma operacao de sistema ligada a um caso de uso da "
        "Secao 2.2 - em particular, todo o eixo principal de mudanca de status e coberto "
        "por UC-07 Atualizar Status do Animal, que aparece como «include» nos "
        "UCs 03, 09 e 10."
    )
    h3(doc, "3.6.1. Resumo dos estados")
    add_table(doc,
        headers=["Estado", "Significado", "Proximo(s) estado(s) tipico(s)"],
        rows=[
            ["RESGATADO",          "Animal recem-resgatado, aguardando triagem veterinaria",                  "EM_TRATAMENTO, OBITO"],
            ["EM_TRATAMENTO",      "Atendimentos veterinarios em andamento (consultas, vacinas, cirurgias)",  "EM_TRATAMENTO (novo atendimento), DISPONIVEL, OBITO"],
            ["DISPONIVEL",         "Apto a adocao, visivel na busca publica (UC-04)",                         "EM_PROCESSO_ADOCAO, OBITO"],
            ["EM_PROCESSO_ADOCAO", "Adocao aprovada, aguardando contrato e entrega",                          "ADOTADO, DISPONIVEL, OBITO"],
            ["ADOTADO",            "Adocao concluida, acompanhamento pos-adocao ativo",                       "DISPONIVEL (devolucao), estado terminal"],
            ["OBITO",              "Animal veio a obito - estado terminal",                                   "-"],
        ],
    )
    h3(doc, "3.6.2. Principais transicoes")
    add_table(doc,
        headers=["De -> Para", "Evento / Operacao", "Caso de Uso"],
        rows=[
            ["RESGATADO -> EM_TRATAMENTO",            "Primeiro registrarAtendimento() apos resgate",       "UC-03 «include» UC-07"],
            ["EM_TRATAMENTO -> EM_TRATAMENTO",        "Novo atendimento (auto-transicao)",                  "UC-03"],
            ["EM_TRATAMENTO -> DISPONIVEL",           "emitirLaudoLiberacao() com liberaParaAdocao = true", "UC-09 «include» UC-07"],
            ["DISPONIVEL -> EM_PROCESSO_ADOCAO",      "aprovarAdocao() pelo Coordenador",                   "UC-10 «include» UC-07"],
            ["EM_PROCESSO_ADOCAO -> ADOTADO",         "concluirAdocao() (contrato assinado)",               "UC-10"],
            ["EM_PROCESSO_ADOCAO -> DISPONIVEL",      "Desistencia do adotante ou rejeicao",                "UC-10"],
            ["ADOTADO -> DISPONIVEL",                 "Devolucao do animal (falha no acompanhamento)",      "UC-06"],
            ["qualquer -> OBITO",                     "Obito clinico do animal",                            "UC-03"],
        ],
    )
    add_image(doc, "25-estado-animal.png", "Figura 25 - Diagrama de Estados: Ciclo de Vida do Animal")


# ----- 4. Modelos de Dados -------------------------------------------------

def build_secao4(doc) -> None:
    h1(doc, "4. Modelos de Dados")
    add_para(doc,
        "Apresentam-se os esquemas de banco de dados do PataAmiga e a estrategia de "
        "mapeamento entre as representacoes de objetos (entidades JPA) e nao-objetos "
        "(tabelas relacionais do PostgreSQL)."
    )
    add_para(doc,
        "O modelo logico espelha o diagrama de classes da Secao 3.3, preservando a "
        "Abordagem C para usuarios: tabela usuario central e cinco tabelas de perfil "
        "opcionais 1:1."
    )

    h2(doc, "4.1. Diagrama Entidade-Relacionamento")
    h3(doc, "4.1.1. Visao geral das tabelas")
    add_table(doc,
        headers=["Tabela", "Cardinalidade com usuario", "Papel"],
        rows=[
            ["usuario",              "-",         "Identidade central; dados pessoais e credenciais"],
            ["perfil_adotante",      "1 - 0..1",  "Dados especificos do papel Adotante"],
            ["perfil_doador",        "1 - 0..1",  "Dados especificos do papel Doador"],
            ["perfil_voluntario",    "1 - 0..1",  "Dados especificos do papel Voluntario"],
            ["perfil_veterinario",   "1 - 0..1",  "Dados especificos do papel Veterinario (CRMV, especialidade)"],
            ["perfil_coordenador",   "1 - 0..1",  "Dados especificos do papel Coordenador"],
            ["animal",               "-",         "Entidade central do negocio"],
            ["registro_veterinario", "-",         "Prontuario; FK para animal e perfil_veterinario"],
            ["adocao",               "-",         "Processo de adocao; FKs para animal, perfil_adotante e perfil_coordenador"],
            ["doacao",               "-",         "Aporte financeiro; FK para perfil_doador"],
            ["relatorio",            "-",         "Saida agregada; FK para perfil_coordenador"],
        ],
    )
    add_para(doc,
        "Nas cinco tabelas perfil_* a coluna usuario_id e simultaneamente PK e FK para "
        "usuario(id) - caracteristica da Abordagem C que garante a cardinalidade 1:1 "
        "opcional sem coluna tipo_perfil discriminadora e sem heranca no banco."
    )
    add_image(doc, "26-entidade-relacionamento.png",
              "Figura 26 - Diagrama Entidade-Relacionamento do PataAmiga (PostgreSQL)")

    h2(doc, "4.2. Estrategia de Mapeamento Objeto-Relacional")
    add_para(doc,
        "O backend usa Hibernate como provider JPA. Cada classe do modelo de dominio "
        "(Secao 3.3) e mapeada para uma tabela PostgreSQL conforme as convencoes a seguir."
    )

    h3(doc, "4.2.1. Convencoes gerais")
    add_table(doc,
        headers=["Aspecto", "Decisao"],
        rows=[
            ["Nomes de tabelas e colunas", "snake_case minusculo (via Hibernate.physical_naming_strategy = CamelCaseToUnderscoresNamingStrategy)"],
            ["Chaves primarias",           "@Id @GeneratedValue(strategy = GenerationType.IDENTITY) mapeado para BIGSERIAL"],
            ["Enums",                      "@Enumerated(EnumType.STRING) - preserva legibilidade no banco e seguranca em refatoracoes"],
            ["Datas/horas",                "LocalDate -> DATE, LocalDateTime -> TIMESTAMP (JPA 3.1, sem @Temporal)"],
            ["Valores monetarios",         "BigDecimal -> NUMERIC(12,2) em Doacao.valor"],
            ["Texto longo",                "Campos como Adocao.acompanhamentoPosAdocao, Animal.descricao e RegistroVeterinario.observacoes usam @Column(columnDefinition = \"TEXT\")"],
            ["JSON estruturado",           "Relatorio.dadosAgregados mapeado como JSONB via @JdbcTypeCode(SqlTypes.JSON) (Hibernate 6)"],
            ["equals / hashCode",          "Baseados apenas no id (entidades JPA com identidade gerada)"],
            ["Versionamento",              "@Version em entidades editaveis com concorrencia (Animal, Adocao) para locking otimista"],
        ],
    )

    h3(doc, "4.2.2. Mapeamento da Abordagem C (Usuario + Perfis)")
    add_para(doc,
        "Cada perfil e uma entidade independente cuja chave primaria e tambem a chave "
        "estrangeira para usuario, materializando o relacionamento 1:1 opcional sem "
        "coluna nullable na tabela base."
    )
    add_code(doc,
        "@Entity\n"
        "@Table(name = \"usuario\")\n"
        "public class Usuario {\n"
        "    @Id @GeneratedValue(strategy = GenerationType.IDENTITY)\n"
        "    private Long id;\n\n"
        "    @Column(nullable = false, unique = true, length = 120)\n"
        "    private String email;\n\n"
        "    @OneToOne(mappedBy = \"usuario\", cascade = CascadeType.ALL,\n"
        "              orphanRemoval = true, fetch = FetchType.LAZY)\n"
        "    private PerfilAdotante perfilAdotante;\n"
        "    // ... mesma anotacao para os outros 4 perfis\n"
        "}\n\n"
        "@Entity\n"
        "@Table(name = \"perfil_adotante\")\n"
        "public class PerfilAdotante {\n"
        "    @Id\n"
        "    private Long id;          // mesmo valor de usuario.id\n\n"
        "    @MapsId\n"
        "    @OneToOne(fetch = FetchType.LAZY)\n"
        "    @JoinColumn(name = \"usuario_id\")\n"
        "    private Usuario usuario;\n"
        "}"
    )
    add_bullets(doc, [
        "@MapsId informa ao Hibernate que a PK do PerfilAdotante e derivada da FK para "
        "Usuario - nao ha sequencia separada.",
        "O lado proprietario do relacionamento e o perfil (ele carrega usuario_id); o "
        "Usuario usa mappedBy.",
        "cascade = ALL + orphanRemoval = true permitem criar/remover o perfil junto com "
        "o usuario quando apropriado.",
    ])
    add_para(doc,
        "Por que nao @Inheritance? Uma hierarquia JPA (SINGLE_TABLE, JOINED ou "
        "TABLE_PER_CLASS) exigiria que cada usuario fosse exatamente um tipo de perfil. "
        "No PataAmiga uma mesma pessoa pode acumular qualquer combinacao de papeis "
        "(ex.: voluntario que tambem adota e doa), o que invalida a heranca e justifica "
        "a Abordagem C com composicao."
    )

    h3(doc, "4.2.3. Mapeamento dos principais relacionamentos")
    add_table(doc,
        headers=["Relacionamento (UML)", "Anotacoes JPA", "Coluna FK no banco"],
        rows=[
            ['Animal "1" -> "0..*" RegistroVeterinario',
             '@OneToMany(mappedBy="animal", cascade=ALL, orphanRemoval=true) em Animal; '
             '@ManyToOne @JoinColumn(name="animal_id") em RegistroVeterinario',
             'registro_veterinario.animal_id'],
            ['PerfilVeterinario "1" -> "0..*" RegistroVeterinario',
             '@ManyToOne(fetch=LAZY) @JoinColumn(name="veterinario_id") em RegistroVeterinario',
             'registro_veterinario.veterinario_id'],
            ['Animal "1" -> "0..*" Adocao',
             '@OneToMany(mappedBy="animal") em Animal; '
             '@ManyToOne @JoinColumn(name="animal_id") em Adocao',
             'adocao.animal_id'],
            ['PerfilAdotante "1" -> "0..*" Adocao',
             '@ManyToOne @JoinColumn(name="adotante_id") em Adocao',
             'adocao.adotante_id'],
            ['PerfilCoordenador "0..1" -> "0..*" Adocao',
             '@ManyToOne(optional=true) @JoinColumn(name="coordenador_id") em Adocao',
             'adocao.coordenador_id (nullable)'],
            ['PerfilDoador "1" -> "0..*" Doacao',
             '@ManyToOne @JoinColumn(name="doador_id") em Doacao',
             'doacao.doador_id'],
            ['PerfilCoordenador "1" -> "0..*" Relatorio',
             '@ManyToOne @JoinColumn(name="coordenador_id") em Relatorio',
             'relatorio.coordenador_id'],
        ],
    )

    h3(doc, "4.2.4. Estrategias de busca (FetchType) e cascata")
    add_bullets(doc, [
        "Default = LAZY em todos os @ManyToOne e @OneToOne - colecoes e referencias so "
        "sao carregadas sob demanda, evitando o N+1. Quando uma view precisa de joins, "
        "a camada Service usa @EntityGraph ou JPQL com JOIN FETCH.",
        "CascadeType.ALL + orphanRemoval = true apenas onde a composicao e exclusiva: "
        "Usuario -> Perfil*, Animal -> RegistroVeterinario. Nas demais relacoes usa-se "
        "CascadeType.PERSIST ou nenhuma cascata, evitando exclusoes em cadeia indesejadas.",
    ])

    h3(doc, "4.2.5. Versionamento de schema")
    add_para(doc,
        "Migrations versionadas com Flyway (db/migration/V<ts>__<descricao>.sql). O "
        "Hibernate roda em modo ddl-auto=validate em producao - o schema e construido "
        "exclusivamente pelas migrations, e a aplicacao apenas valida na inicializacao. "
        "Isso garante que o ER documentado nesta secao corresponde ao banco real."
    )


# ============================================================================
#  Entry point
# ============================================================================

def main() -> None:
    template = find_template()
    print(f"Template:  {template}")
    print(f"Imagens:   {IMG_DIR}")
    print(f"Saida:     {OUTPUT}")

    doc = Document(str(template))
    clear_body(doc)

    build_cover(doc)
    build_historico(doc)
    build_toc(doc)
    build_intro(doc)
    build_secao2(doc)
    build_secao3(doc)
    build_secao4(doc)

    doc.save(str(OUTPUT))
    size_kb = OUTPUT.stat().st_size / 1024
    print(f"OK: {OUTPUT.name} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
